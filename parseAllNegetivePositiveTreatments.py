# The files in the process folder under files -> <State> -> Process
# are only the negetive treatments.

# We want to calculate case load, which includes all types of treatments
# These files can be found under  files -> <State> -> DOCX



# ! This peice of code helped understand the docx file structure.
# from docx import Document

# def print_paragraphs_xml(file_path):
#     doc = Document(file_path)
#     for paragraph in doc.paragraphs:
#         print(paragraph._element.xml)


# file_path = "./files/Alabama/DOCX/Westlaw Precision - List of 10 results for adv CO(HIGH) And DA(AFTER 12311824) And DA(BEFORE 111826).docx"  # Replace this with the path to your .docx file
# print_paragraphs_xml(file_path)

from docx import Document
from xml.etree import ElementTree as ET
import re
import os
import csv


def extract_text_from_paragraphs(file_path):
    doc = Document(file_path)
    case_count = 0
    for paragraph in doc.element.body:
        xml_str = ET.tostring(paragraph, encoding='unicode', method='xml')
        paragraph_xml = ET.fromstring(xml_str)
        for run in paragraph_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
            for text_element in run.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                text = text_element.text.strip()
                if text:
                    # print(text)
                    if re.match(r"\d+\.", text):
                        print(text)
                        print("$$$$$$$$$$$$$$$$$$$$")
                        # ! This is to remove edge cases and wrong counting of text such as 1942. somewhere in the paragraph etc.
                        try:
                          if int(text.replace(".", "")) == case_count + 1:
                            case_count += 1
                        except:
                            pass
    print("===============================")
    print(case_count)
    print("===============================")
    return case_count

stateFolderNames = [
    "Alabama",
    "Connecticut",
    "Illinois",
    "Maine",
    "Missouri", 
    "NewMexico", 
    "Oklahomasup",
    "Tennessee",
    "Washington",
    "Alaska",
    "Delaware",
    "Indiana", 
    "Maryland",
    "Montana",
    "NewYork",
    "Oregon",
    "Texascri",
    "WestVirginia",
    "Arizona",
    "Florida",
    "Iowa",
    "Massachusetts",
    "Nebraska",
    "NorthCarolina",
    "Pennsylvania",
    "Texassup",
    "Wisconsin",
    "Arkansas",
    "Georgia",
    "Kansas",
    "Michigan",
    "Nevada",
    "NorthDakota",
    "Rhode",
    "Utah",
    "Wyoming",
    "California",
    "Hawaii",
    "Kentucky",
    "Minnesota",
    "NewHampshire",
    "Ohio",
    "SouthCarolina",
    "Vermont",
    "Colorado",
    "Idaho",
    "Louisiana",
    "Mississippi",
    "NewJersey",
    "Oklahomacri",
    "SouthDakota",
    "Virgina",								
]

final_output_file_name = "./outputAllTreatmentFromDocx.csv"

with open(final_output_file_name, 'a') as f:
    # using csv.writer method from CSV package
    write = csv.writer(f)
    write.writerow(["STATE", "CASE_LOAD"])


for state in stateFolderNames:
    # Define the folder containing the DOCX files for a perticular State
    folder_path = f"./files/{state}/DOCX" #'/files/<STATE>/Process'

    # print(os.listdir(folder_path)) #DEBUG

    state_case_load = 0

    # Iterate over all files in the folder
    for file_name in os.listdir(folder_path):
        print(file_name)
        if file_name.endswith('.docx'):
            # Construct the absolute path to the file
            docx_file_path = os.path.join(folder_path, file_name)
            state_case_load += extract_text_from_paragraphs(docx_file_path)

    print("*******************")
    print(state, state_case_load)
    print("*******************")

    final_output_file_name = "./outputAllTreatmentFromDocx.csv"

    with open(final_output_file_name, 'a') as f:
        # using csv.writer method from CSV package
        write = csv.writer(f)
        write.writerow([state, state_case_load])